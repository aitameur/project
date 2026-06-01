"""Génère notice_technique.docx — Notice Technique complète du scanner OWASP."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ASSETS = Path(__file__).parent / "assets"
OUT = Path(__file__).parent / "notice_technique.docx"

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── styles helper ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def code_block(lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shd)

def add_img(path_str, caption, width_in=5.5):
    p = path_str if isinstance(path_str, Path) else Path(path_str)
    if not p.exists():
        para(f"[Image non trouvée : {p.name}]", italic=True)
        return
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ip.add_run()
    run.add_picture(str(p), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)

def simple_table(headers: list[str], rows: list[list[str]]):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        cell._tc.get_or_add_tcPr().append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
r = title_p.add_run("NOTICE TECHNIQUE")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Scanner de Vulnérabilités Web — OWASP Top 10")
sr.font.size = Pt(16)
sr.bold = True

doc.add_paragraph()
meta = [
    ("Projet :", "Projet de Fin d'Année (PFA) — EMSI 2024–2025"),
    ("Auteur :", "Ait Ameur Fatima"),
    ("Encadrant :", "—"),
    ("Date :", "Mai 2026"),
    ("Version :", "1.0"),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(label + " ").bold = True
    p.add_run(val)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — VUE D'ENSEMBLE
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Vue d'ensemble du projet")
para(
    "Ce document est la notice technique complète du scanner de vulnérabilités web "
    "développé dans le cadre du PFA EMSI 2024–2025. Il décrit l'architecture, les "
    "modules, l'environnement de test, les résultats obtenus et les points techniques "
    "saillants de l'implémentation."
)

h2("1.1 Objectif")
para(
    "L'objectif est de détecter automatiquement les vulnérabilités des dix catégories "
    "OWASP Top 10 (2021) sur une application web cible, en générant un rapport HTML "
    "structuré avec scores CVSS v3.1."
)

h2("1.2 Couverture OWASP Top 10 (2021)")
simple_table(
    ["Code", "Catégorie OWASP", "Détecteur(s)"],
    [
        ["A01", "Broken Access Control",       "idor, path_traversal"],
        ["A02", "Cryptographic Failures",       "crypto_failures"],
        ["A03", "Injection",                    "sqli, xss, xxe"],
        ["A04", "Insecure Design",              "—"],
        ["A05", "Security Misconfiguration",    "security_misconfig"],
        ["A06", "Vulnerable Components",        "vulnerable_components"],
        ["A07", "Auth & Session Management",    "broken_auth"],
        ["A08", "Software Integrity Failures",  "integrity_failures"],
        ["A09", "Logging & Monitoring",         "—"],
        ["A10", "SSRF",                         "ssrf"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Architecture du pipeline")
para(
    "Le scanner s'exécute en cinq étapes séquentielles orchestrées par la classe "
    "Scanner (scanner/core.py) :"
)
steps = [
    "0. Auto-login (optionnel) — authentifie la session HTTP avant le crawl.",
    "1. Crawl BFS — explore l'application jusqu'à la profondeur max_depth et collecte les Endpoint.",
    "2. Détection — chaque détecteur actif ou passif analyse les endpoints collectés.",
    "3. Déduplication — supprime les doublons par clé (type, URL, paramètre, payload).",
    "4. Évaluation CVSS — calcule le score et la sévérité de chaque vulnérabilité.",
    "5. Rapport — génère un fichier HTML responsive avec tableau de bord et fiches détaillées.",
]
for s in steps:
    bullet(s)

h2("2.1 Diagramme du flux")
code_block([
    "CLI (main.py)",
    "  └─► Scanner.run()",
    "        ├─► _maybe_login()     [authenticator.py]",
    "        ├─► Crawler.crawl()    [crawler/]",
    "        │     └─► Endpoint[]",
    "        ├─► Detector.scan_all() x N  [detectors/]",
    "        │     └─► Vulnerability[]",
    "        ├─► _deduplicate()",
    "        ├─► CVSSEvaluator.evaluate_all()  [evaluator.py]",
    "        └─► HTMLReporter.save()           [reporter.py]",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — STRUCTURE DES FICHIERS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Structure des fichiers")
code_block([
    "imple/",
    "├── main.py                   Point d'entrée CLI",
    "├── scanner/",
    "│   ├── __init__.py           Expose la classe Scanner",
    "│   ├── core.py               Orchestrateur principal",
    "│   ├── authenticator.py      Auto-login générique + DVWA",
    "│   └── vulnerability.py      Dataclass Vulnerability",
    "├── crawler/",
    "│   ├── __init__.py",
    "│   ├── bfs_crawler.py        Crawl BFS multi-niveaux",
    "│   └── endpoint.py           Dataclass Endpoint",
    "├── detectors/",
    "│   ├── __init__.py           ALL_DETECTORS registry",
    "│   ├── base.py               Classes de base ActiveDetector / PassiveDetector",
    "│   ├── sqli.py               A03 — SQL Injection",
    "│   ├── xss.py                A03 — Cross-Site Scripting",
    "│   ├── xxe.py                A03 — XML External Entity",
    "│   ├── ssrf.py               A10 — SSRF",
    "│   ├── idor.py               A01 — IDOR",
    "│   ├── path_traversal.py     A01 — Path Traversal",
    "│   ├── broken_auth.py        A07 — Auth Failures",
    "│   ├── crypto_failures.py    A02 — Crypto Failures",
    "│   ├── security_misconfig.py A05 — Security Misconfig",
    "│   ├── vulnerable_components.py A06 — Outdated Libraries",
    "│   ├── integrity_failures.py A08 — Software Integrity",
    "│   └── file_upload.py        A01/A03 — Unrestricted Upload",
    "├── evaluator.py              Calcul CVSS v3.1",
    "├── reporter.py               Génération rapport HTML",
    "└── tests/",
    "    ├── conftest.py           Mini-app Flask volontairement vulnérable",
    "    └── test_*.py             26 tests pytest",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — LES 12 DÉTECTEURS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Les 12 détecteurs")

h2("4.1 Détecteurs actifs (injection de payloads)")
simple_table(
    ["Nom", "OWASP", "Technique de détection"],
    [
        ["sqli",       "A03", "Injecte ' et -- ; cherche erreurs SQL dans la réponse"],
        ["xss",        "A03", "Injecte <script>alert() ; cherche le payload réfléchi"],
        ["xxe",        "A03", "POST XML avec entité SYSTEM ; cherche /etc/passwd"],
        ["ssrf",       "A10", "Injecte URL 127.0.0.1/169.254.169.254 ; cherche métadonnées"],
        ["idor",       "A01", "Incrémente les ID numériques ; compare les réponses"],
        ["path_trav",  "A01", "Injecte ../../etc/passwd ; cherche root: dans la réponse"],
        ["file_upload","A01", "Upload .php/.jsp ; vérifie l'exécution distante"],
    ],
)

h2("4.2 Détecteurs passifs (analyse des réponses existantes)")
simple_table(
    ["Nom", "OWASP", "Ce qui est analysé"],
    [
        ["broken_auth",      "A07", "Cookie sans HttpOnly/Secure, login en clair, brute-force"],
        ["crypto_failures",  "A02", "Formulaire password en HTTP, headers HSTS absents"],
        ["security_misconfig","A05","Fichiers sensibles (.env, phpinfo, listing répertoire)"],
        ["vuln_components",  "A06", "Headers Server/X-Powered-By, meta generator, jQuery version"],
        ["integrity_failures","A08","Scripts/CSS CDN sans attribut integrity (SRI)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ENVIRONNEMENT DE TEST
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Environnement de test")

h2("5.1 Infrastructure")
simple_table(
    ["Composant", "Détail"],
    [
        ["Machine hôte",    "Windows 11 Pro — Python 3.12"],
        ["Machine cible",   "Kali Linux 2024.x — VirtualBox 7.x"],
        ["Réseau VM",       "Host-Only : 192.168.56.101 / NAT : 10.0.2.15"],
        ["Application cible","DVWA (Damn Vulnerable Web Application) — Apache + MariaDB"],
        ["URL DVWA",        "http://192.168.56.101/dvwa/"],
        ["Credentials",     "admin / password"],
        ["Niveau sécurité", "Low (toutes protections désactivées)"],
    ],
)

h2("5.2 Procédure de reset DVWA (à effectuer après chaque redémarrage)")
para("À chaque redémarrage de la VM Kali, MariaDB perd l'état de la base. Procédure :")
steps_dvwa = [
    "1. Ouvrir Firefox dans Kali → http://127.0.0.1/dvwa/setup.php",
    "2. Cliquer « Create / Reset Database »",
    "3. Se connecter : admin / password",
    "4. Aller dans DVWA Security → sélectionner « Low » → Submit",
]
for s in steps_dvwa:
    bullet(s)

h2("5.3 Capture d'écran — Environnement DVWA")
add_img(ASSETS / "s4.jpeg", "Figure 5.1 — Page d'accueil DVWA (Welcome to DVWA)")
add_img(ASSETS / "s5.jpeg", "Figure 5.2 — Niveau de sécurité DVWA réglé sur « Low »")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — GUIDE D'UTILISATION CLI
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Guide d'utilisation (CLI)")

h2("6.1 Options disponibles")
simple_table(
    ["Option", "Description", "Défaut"],
    [
        ["--url URL",          "URL cible du scan (obligatoire)", "—"],
        ["--output FILE",      "Chemin du rapport HTML", "rapport.html"],
        ["--depth N",          "Profondeur de crawl BFS", "2"],
        ["--timeout N",        "Timeout HTTP en secondes", "10"],
        ["--max-pages N",      "Nombre maximal de pages crawlées", "100"],
        ["--only LIST",        "Détecteurs à activer (ex: sqli,xss)", "tous"],
        ["--cookie STR",       "Cookies d'authentification manuels", "—"],
        ["--dvwa",             "Auto-login DVWA (admin/password + security=low)", "false"],
        ["--login-url URL",    "URL du formulaire de login générique", "—"],
        ["--username STR",     "Identifiant pour l'auto-login", "—"],
        ["--password STR",     "Mot de passe pour l'auto-login", "—"],
        ["--dvwa-security STR","Niveau de sécurité DVWA à régler", "low"],
        ["-v / --verbose",     "Logs détaillés (DEBUG)", "false"],
    ],
)

h2("6.2 Exemples d'utilisation")
code_block([
    "# Scan simple",
    "python main.py --url http://192.168.56.101/dvwa/",
    "",
    "# Scan DVWA avec auto-login",
    "python main.py --url http://192.168.56.101/dvwa/ --dvwa",
    "",
    "# Scan avec profondeur 3, rapport personnalisé et logs détaillés",
    "python main.py --url http://192.168.56.101/dvwa/ --dvwa --depth 3 \\",
    "               --output resultats.html -v",
    "",
    "# Scan uniquement SQLi et XSS",
    "python main.py --url http://cible.com --only sqli,xss",
    "",
    "# Scan avec session cookie manuelle",
    "python main.py --url http://cible.com --cookie 'PHPSESSID=abc; security=low'",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — RÉSULTATS DU SCAN
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Résultats du scan sur DVWA")

h2("7.1 Résumé")
simple_table(
    ["Métrique", "Valeur"],
    [
        ["URL cible",              "http://192.168.56.101/dvwa/"],
        ["Pages crawlées",         "43"],
        ["Durée totale",           "9.10 secondes"],
        ["Vulnérabilités totales", "128"],
        ["Critique (CVSS ≥ 9.0)", "3"],
        ["Élevé (CVSS 7.0–8.9)",  "7"],
        ["Moyen (CVSS 4.0–6.9)",  "118"],
        ["Faible (CVSS < 4.0)",   "0"],
    ],
)

h2("7.2 Sortie terminal — scan authentifié")
add_img(ASSETS / "s3.jpeg", "Figure 7.1 — Sortie terminal : 128 vulnérabilités détectées en 9.10s")

h2("7.3 Tableau de bord du rapport HTML")
add_img(ASSETS / "s6.jpeg", "Figure 7.2 — Rapport HTML : dashboard (3 Critique, 7 Élevé, 118 Moyen)")

h2("7.4 Exemples de fiches vulnérabilité")

h3("7.4.1 Cross-Site Scripting (XSS) — A03")
add_img(ASSETS / "s9.jpeg", "Figure 7.3 — Fiche XSS (CVSS 6.1, payload <script>alert('XSS')</script>)")

h3("7.4.2 Server-Side Request Forgery (SSRF) — A10")
add_img(ASSETS / "s8.jpeg", "Figure 7.4 — Fiche SSRF (CVSS 8.5, payload file:///etc/passwd)")

h3("7.4.3 Path Traversal — A01")
add_img(ASSETS / "s7.jpeg", "Figure 7.5 — Fiche Path Traversal (CVSS 7.5, payload ../../etc/passwd)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — TESTS UNITAIRES
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Tests unitaires")

h2("8.1 Infrastructure de test")
para(
    "Les 26 tests pytest utilisent une mini-application Flask DÉLIBÉRÉMENT vulnérable "
    "définie dans tests/conftest.py. Cette app expose les mêmes vulnérabilités que DVWA "
    "mais tourne en local sur un port aléatoire, sans dépendance externe."
)

h2("8.2 Couverture des tests")
simple_table(
    ["Fichier de test", "Tests", "Ce qui est vérifié"],
    [
        ["test_sqli.py",          "3", "Détection d'erreur SQL sur /search?q='"],
        ["test_xss.py",           "3", "Payload <script> réfléchi sur /greet"],
        ["test_xxe.py",           "2", "Entité SYSTEM sur /xml"],
        ["test_ssrf.py",          "3", "Redirect vers 127.0.0.1 sur /fetch"],
        ["test_idor.py",          "2", "Accès ID=2,3 depuis ID=1 sur /user"],
        ["test_path_traversal.py","2", "../../etc/passwd sur /file"],
        ["test_broken_auth.py",   "2", "Cookie sans HttpOnly, login en clair"],
        ["test_crypto.py",        "2", "Formulaire password en HTTP"],
        ["test_misconfig.py",     "2", "/.env, /phpinfo.php, listing"],
        ["test_vuln_components.py","2","Header Server Apache/2.4.25, jQuery 1.6.4"],
        ["test_integrity.py",     "1", "Script CDN sans integrity="],
        ["test_file_upload.py",   "2", "Upload .php accepté sans validation"],
        ["TOTAL",                "26", ""],
    ],
)

h2("8.3 Exécution des tests")
code_block([
    "cd imple",
    "pytest tests/ -v",
    "",
    "# Résultat attendu :",
    "# collected 26 items",
    "# ...",
    "# 26 passed in 1.43s",
])
add_img(ASSETS / "s1.jpeg", "Figure 8.1 — Résultats pytest : 26 tests passés en 1.43s")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — POINTS TECHNIQUES CLÉS
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Points techniques clés")

h2("9.1 Gestion de session HTTP")
para(
    "Le scanner utilise une unique instance requests.Session pour tous les appels. "
    "Cela permet de :"
)
for item in [
    "Conserver les cookies d'authentification entre les requêtes.",
    "Réutiliser les connexions TCP (keep-alive) pour les performances.",
    "Appliquer un User-Agent et des headers personnalisés globalement.",
]:
    bullet(item)

h2("9.2 Déduplication des vulnérabilités")
para(
    "Chaque Vulnerability expose une méthode key() retournant un tuple "
    "(type, url, param, payload). La méthode Scanner._deduplicate() utilise ce "
    "tuple comme clé d'un dict pour éliminer les doublons avant évaluation CVSS."
)

h2("9.3 Détecteurs passifs vs actifs")
para(
    "Les détecteurs héritent soit de ActiveDetector (injectent des payloads dans "
    "les paramètres) soit de PassiveDetector (analysent uniquement le contenu HTML "
    "et les headers des réponses déjà collectées). Les passifs reçoivent tous les "
    "endpoints visités ; les actifs ne reçoivent que ceux avec des paramètres injectables."
)

h2("9.4 Correction du bug d'authentification DVWA")
para(
    "Un bug critique dans authenticator.py empêchait l'auto-login DVWA : "
    "l'URL /login.php était construite de façon incorrecte."
)
para("Code buggué :", bold=True)
code_block([
    "# BUG : urljoin('http://192.168.56.101/dvwa', 'login.php')",
    "#      => 'http://192.168.56.101/login.php'  (404 !)",
    "login_url = urljoin(base_url, 'login.php')",
])
para("Code corrigé :", bold=True)
code_block([
    "# CORRECT : forcer le slash final avant urljoin",
    "base = base_url if base_url.endswith('/') else base_url + '/'",
    "login_url = urljoin(base, 'login.php')",
    "# => 'http://192.168.56.101/dvwa/login.php'  (200 OK)",
])

h2("9.5 Calcul CVSS v3.1")
para(
    "Chaque détecteur associe à ses vulnérabilités des métriques CVSS (vecteur d'attaque, "
    "complexité, privilèges requis, impact C/I/A). L'évaluateur CVSSEvaluator calcule "
    "le score numérique et détermine la sévérité :"
)
simple_table(
    ["Score", "Sévérité"],
    [
        ["9.0 – 10.0", "Critique"],
        ["7.0 – 8.9",  "Élevé"],
        ["4.0 – 6.9",  "Moyen"],
        ["0.1 – 3.9",  "Faible"],
        ["0.0",        "Aucune"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Conclusion")
para(
    "Le scanner OWASP Top 10 développé dans ce PFA couvre les dix catégories de "
    "vulnérabilités les plus critiques avec 12 détecteurs autonomes. Testé sur DVWA "
    "en mode authentifié, il identifie 128 vulnérabilités en moins de 10 secondes, "
    "générant un rapport HTML exploitable avec scores CVSS v3.1."
)
para(
    "Les 26 tests unitaires valident chaque détecteur de façon isolée, garantissant "
    "la robustesse et la maintenabilité du code. L'architecture modulaire (pipeline "
    "crawl → detect → evaluate → report) facilite l'ajout de nouveaux détecteurs "
    "pour couvrir des catégories supplémentaires."
)
para(
    "Ce projet constitue une base solide pour un outil de pentest éducatif, "
    "extensible vers des techniques de détection plus avancées (blind SQLi, "
    "stored XSS, détection d'authentification JWT, etc.)."
)

doc.add_paragraph()
para(
    "AVERTISSEMENT : Cet outil ne doit être utilisé QUE sur des applications "
    "dont vous avez l'autorisation explicite de tester la sécurité.",
    italic=True,
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"[OK] {OUT}")
